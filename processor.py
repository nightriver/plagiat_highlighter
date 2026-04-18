import re
import io
import math
from difflib import SequenceMatcher
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from rapidfuzz.distance import Levenshtein


PAGE_MARKER_RE = re.compile(r'^\s*[Сс]\.\s*\d+[\d\s\-–—]*\s*$')
DASH_MARKER_RE = re.compile(r'^--$')


def is_page_marker(text: str) -> bool:
    t = text.strip()
    return bool(PAGE_MARKER_RE.match(t) or DASH_MARKER_RE.match(t))


def normalize_hyphenated(text: str) -> str:
    """
    'Ін- тернет' → 'Інтернет'. прибирає переноси у словах.
    """
    return re.sub(r'(\w)-\s+(\w)', r'\1\2', text)


COMMENT_RE = re.compile(r'^(\d+\.|Посилання на:|Там само$)')
URL_RE = re.compile(r'^https?://')


# ---------------------------------------------------------------------------
# Нова логіка: повертають ANNOTATED список всіх параграфів ячейки
# кожен елемент: {'para': para, 'zone': 'text'|'plain'}
# 'text'  = ці параграфи будуть підсвічуватись
# 'plain' = ці параграфи переписуються без змін
# ---------------------------------------------------------------------------

def annotate_left_cell(cell) -> tuple:
    """
    Ліва ячейка.
    Зони:
      - до PAGE MARKER (включно) → 'plain'
      - параграфи після PAGE MARKER до першого порожнього → 'text'
      - після порожнього (comment-зона) → 'plain'
    """
    paragraphs = cell.paragraphs
    marker_idx = None
    for i, para in enumerate(paragraphs):
        if is_page_marker(para.text):
            marker_idx = i
            break

    if marker_idx is None:
        annotated = [{'para': p, 'zone': 'text'} for p in paragraphs if p.text.strip()]
        return annotated, "WARNING: page marker not found"

    annotated = []
    # До маркера + сам маркер → plain
    for i in range(marker_idx + 1):
        annotated.append({'para': paragraphs[i], 'zone': 'plain'})

    # Після маркера: текст до першого порожнього
    in_text = True
    for para in paragraphs[marker_idx + 1:]:
        if in_text and para.text.strip() == '':
            in_text = False
            annotated.append({'para': para, 'zone': 'plain'})
            continue
        if in_text:
            annotated.append({'para': para, 'zone': 'text'})
        else:
            annotated.append({'para': para, 'zone': 'plain'})

    return annotated, None


def annotate_right_cell(cell) -> tuple:
    """
    Права ячейка.
    Зони:
      - до ОСТАННЬОГО PAGE MARKER (включно) → 'plain'
      - після аж до COMMENT_RE / кінця ячейки → 'text'
    """
    paragraphs = cell.paragraphs
    marker_idx = None
    for i, para in enumerate(paragraphs):
        if is_page_marker(para.text):
            marker_idx = i  # не break — шукаємо останній

    if marker_idx is None:
        annotated = [{'para': p, 'zone': 'text'} for p in paragraphs if p.text.strip()]
        return annotated, "WARNING: page marker not found"

    annotated = []
    # До останнього маркера + сам маркер → plain
    for i in range(marker_idx + 1):
        annotated.append({'para': paragraphs[i], 'zone': 'plain'})

    # Після маркера: text до COMMENT_RE
    for para in paragraphs[marker_idx + 1:]:
        txt = para.text.strip()
        if not txt:
            annotated.append({'para': para, 'zone': 'plain'})
            continue
        if COMMENT_RE.match(txt):
            # Коментар і все після → plain
            annotated.append({'para': para, 'zone': 'plain'})
            # решта параграфів цього рядку теж plain
            break
        annotated.append({'para': para, 'zone': 'text'})

    return annotated, None


# ---------------------------------------------------------------------------
# Вертикальне вирівнювання: додає порожні рядки на початок лівої ячейки,
# щоб текст починався на тій самій висоті, що й текст правої.
# Використовує оцінку візуальних рядків — враховує перенос тексту в колонці.
# ---------------------------------------------------------------------------

class _EmptyPara:
    """Фіктивний параграф з порожнім текстом — лише для padding."""
    text = ''


# Коефіцієнт середньої ширини символу відносно розміру шрифту.
# Calibri/Arial: кирилиця трохи ширша за латиницю, ~0.55 добре апроксимує.
_CHAR_WIDTH_FACTOR = 0.55


def _estimate_visual_lines(text: str, chars_per_line: float) -> int:
    """
    Оцінює кількість візуальних рядків, які займає параграф.
    Порожній параграф = 1 рядок (займає висоту навіть якщо пустий).
    """
    if not text.strip():
        return 1
    return max(1, math.ceil(len(text) / chars_per_line))


def _count_leading_visual_lines(annotated: list, chars_per_line: float) -> int:
    """
    Підраховує суму візуальних рядків у 'plain'-зоні до першого 'text'.
    """
    total = 0
    for e in annotated:
        if e['zone'] == 'text':
            break
        total += _estimate_visual_lines(e['para'].text, chars_per_line)
    return total


def pad_left_to_match_right(
    left_annotated: list,
    right_annotated: list,
    font_size: int,
    right_cell_width_emu: int,
) -> list:
    """
    Вставляє порожні рядки на початок лівої ячейки так, щоб перший 'text'-параграф
    починався приблизно на тій самій висоті, що й перший 'text'-параграф правої.

    right_cell_width_emu — ширина правої колонки в EMU (English Metric Units).
    1 pt = 12700 EMU. Якщо ширина невідома — використовується fallback 400 pt.
    """
    cell_width_pt  = (right_cell_width_emu / 12700) if right_cell_width_emu else 400.0
    avg_char_pt    = font_size * _CHAR_WIDTH_FACTOR
    chars_per_line = max(1.0, cell_width_pt / avg_char_pt)

    right_lines = _count_leading_visual_lines(right_annotated, chars_per_line)
    left_lines  = _count_leading_visual_lines(left_annotated,  chars_per_line)
    padding     = right_lines - left_lines

    if padding <= 0:
        return left_annotated

    pad_entries = [{'para': _EmptyPara(), 'zone': 'plain'} for _ in range(padding)]
    return pad_entries + left_annotated


# ---------------------------------------------------------------------------
# Токенізація / порівняння
# ---------------------------------------------------------------------------

def tokenize_paragraphs(paragraphs: list) -> list:
    result = []
    for idx, para in enumerate(paragraphs):
        text = normalize_hyphenated(para.text)
        tokens = re.split(r'(\s+)', text)
        for t in tokens:
            if t:
                result.append((t, idx))
        if idx < len(paragraphs) - 1:
            result.append(('\n', idx))
    return result


def is_word(token: str) -> bool:
    return bool(re.search(r'\w', token))


def fuzzy_match(a: str, b: str, threshold: float) -> bool:
    if not a or not b:
        return False
    min_len = min(len(a), len(b))
    effective_threshold = 70.0 if min_len <= 4 else threshold
    ratio = Levenshtein.normalized_similarity(a, b) * 100
    return ratio >= effective_threshold


def align_tokens(left_pairs: list, right_pairs: list, threshold: float = 75.0) -> tuple:
    left_words  = [(t, i) for t, i in left_pairs  if is_word(t)]
    right_words = [(t, i) for t, i in right_pairs if is_word(t)]

    left_keys  = [re.sub(r'[^\w]', '', t.lower()) for t, _ in left_words]
    right_keys = [re.sub(r'[^\w]', '', t.lower()) for t, _ in right_words]

    matcher = SequenceMatcher(None, left_keys, right_keys, autojunk=False)

    left_status  = {i: 'diff' for i in range(len(left_words))}
    right_status = {i: 'diff' for i in range(len(right_words))}

    for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
        if opcode == 'equal':
            for i in range(i1, i2): left_status[i]  = 'match'
            for j in range(j1, j2): right_status[j] = 'match'
        elif opcode == 'replace':
            right_matched = set()
            for i in range(i1, i2):
                best_j = None
                best_ratio = -1.0
                for j in range(j1, j2):
                    if j in right_matched:
                        continue
                    a, b = left_keys[i], right_keys[j]
                    if not a or not b:
                        continue
                    min_len = min(len(a), len(b))
                    eff_threshold = 70.0 if min_len <= 4 else threshold
                    ratio = Levenshtein.normalized_similarity(a, b) * 100
                    if ratio >= eff_threshold and ratio > best_ratio:
                        best_ratio = ratio
                        best_j = j
                if best_j is not None:
                    left_status[i]  = 'match'
                    right_status[best_j] = 'match'
                    right_matched.add(best_j)

    def rebuild(token_pairs, words, statuses):
        word_iter = iter(enumerate(words))
        current = next(word_iter, None)
        result = []
        for token, para_idx in token_pairs:
            if is_word(token):
                w_idx = current[0] if current is not None else None
                status = statuses.get(w_idx, 'diff') if w_idx is not None else 'diff'
                result.append((token, para_idx, status))
                current = next(word_iter, None)
            else:
                result.append((token, para_idx, None))
        return result

    left_res  = rebuild(left_pairs,  left_words,  left_status)
    right_res = rebuild(right_pairs, right_words, right_status)

    def fill_space_status(token_list):
        result = list(token_list)
        for i in range(1, len(result) - 1):
            token, para_idx, status = result[i]
            if status is None and token != '\n':
                prev_status = result[i - 1][2]
                next_status = result[i + 1][2]
                if prev_status is not None and prev_status == next_status:
                    result[i] = (token, para_idx, prev_status)
        return result

    return fill_space_status(left_res), fill_space_status(right_res)


# ---------------------------------------------------------------------------
# Запис в ячейку
# ---------------------------------------------------------------------------

HIGHLIGHT = {
    'match': WD_COLOR_INDEX.YELLOW,
    'diff':  WD_COLOR_INDEX.TURQUOISE,
}


def rewrite_cell(cell, annotated: list, token_result: list,
                 font_name: str, font_size: int):
    """
    annotated  — весь список параграфів з зонами {'para':..., 'zone': 'plain'|'text'}
    token_result — вже підсвічені токени тільки для 'text'-зони: [(token, para_idx, status)]

    ВАЖЛИВО: тексти plain-параграфів зберігаємо ДО будь-яких змін у DOM,
    бо annotated[0]['para'] і first_para — це один і той самий Python-об'єкт.
    """
    # ── 1. Зберегти тексти plain-параграфів ДО будь-яких змін ──────────────
    saved_texts = [entry['para'].text for entry in annotated]

    # ── 2. Очистити ячейку (залишити лише перший <w:p>) ─────────────────────
    tc = cell._tc
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for p in tc.findall(f'{{{ns}}}p')[1:]:
        tc.remove(p)
    first_para = cell.paragraphs[0]
    first_para.clear()
    current_para = first_para
    first_used = False

    # ── 3. Побудувати карту токенів для text-зони ────────────────────────────
    text_para_map = {}
    for token, para_idx, status in token_result:
        if token == '\n':
            continue
        text_para_map.setdefault(para_idx, []).append((token, status))

    text_para_counter = 0

    # ── 4. Записати всі параграфи ────────────────────────────────────────────
    for entry_idx, entry in enumerate(annotated):
        zone = entry['zone']

        if not first_used:
            p = current_para
            first_used = True
        else:
            p = cell.add_paragraph()
            current_para = p

        if zone == 'plain':
            run = p.add_run(saved_texts[entry_idx])
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = False
            run.italic = False
        else:  # 'text'
            tokens_for_para = text_para_map.get(text_para_counter, [])
            text_para_counter += 1
            for token, status in tokens_for_para:
                run = p.add_run(token)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.bold = False
                run.italic = False
                run.font.highlight_color = HIGHLIGHT.get(status)


# ---------------------------------------------------------------------------
# Головна функція
# ---------------------------------------------------------------------------

def process_document(
    docx_bytes: bytes,
    threshold: float,
    font_name: str,
    font_size: int,
    progress_callback=None
) -> tuple:
    warnings = []
    stats = {'processed': 0, 'warnings': 0, 'skipped': 0}

    doc = Document(io.BytesIO(docx_bytes))
    table = doc.tables[0]
    total_rows = len(table.rows)

    for row_idx, row in enumerate(table.rows):
        if progress_callback:
            progress_callback(row_idx, total_rows, row_idx + 1)

        if len(row.cells) < 2:
            stats['skipped'] += 1
            continue

        left_cell, right_cell = row.cells[0], row.cells[1]

        if not left_cell.text.strip() and not right_cell.text.strip():
            stats['skipped'] += 1
            continue

        left_annotated,  warn_l = annotate_left_cell(left_cell)
        right_annotated, warn_r = annotate_right_cell(right_cell)

        for warn, side in [(warn_l, 'ліва'), (warn_r, 'права')]:
            if warn:
                warnings.append(f"Рядок {row_idx + 1} ({side} колонка): {warn}")
                stats['warnings'] += 1

        # ── Вирівнювання: лівій ячейці додаємо порожні рядки зверху ─────────
        right_width_emu = right_cell.width or 0
        left_annotated = pad_left_to_match_right(
            left_annotated, right_annotated, font_size, right_width_emu
        )

        left_text_paras  = [e['para'] for e in left_annotated  if e['zone'] == 'text']
        right_text_paras = [e['para'] for e in right_annotated if e['zone'] == 'text']

        if not left_text_paras or not right_text_paras:
            stats['skipped'] += 1
            continue

        left_pairs  = tokenize_paragraphs(left_text_paras)
        right_pairs = tokenize_paragraphs(right_text_paras)

        left_result, right_result = align_tokens(left_pairs, right_pairs, threshold)

        rewrite_cell(left_cell,  left_annotated,  left_result,  font_name, font_size)
        rewrite_cell(right_cell, right_annotated, right_result, font_name, font_size)

        stats['processed'] += 1

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read(), warnings, stats
